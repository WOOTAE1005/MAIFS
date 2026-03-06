"""
KIPS 학술발표대회 논문 생성기
DAAC 논문을 KIPS 양식에 맞춰 Word 문서로 생성
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '425')
    sectPr.append(cols)


def set_font(run, size, bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = '맑은 고딕'
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.insert(0, rFonts)


def add_para(doc, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             sb=0, sa=2, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if indent is not None:
        p.paragraph_format.first_line_indent = Pt(indent)
    r = p.add_run(text)
    set_font(r, size, bold=bold)
    return p


def add_heading(doc, text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size, bold=True)
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 9, bold=True)
    return p


def add_table(doc, headers, rows, size=8):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size, bold=True)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_font(r, size)
    return tbl


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 8, bold=True)
    return p


def add_ref(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    r = p.add_run(text)
    set_font(r, 8)
    return p


def main():
    doc = Document()

    # 페이지 여백 (A4)
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # ── 1단: 제목·저자·초록 ──────────────────────────────────────

    # 한국어 제목
    add_para(doc,
             "DAAC: 다중 에이전트 이미지 위변조 탐지를 위한 불일치 인식 적응형 합의",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)

    # 한국어 저자/소속
    add_para(doc, "정준1, 문수혁2",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    add_para(doc, "1동서울대학교 컴퓨터소프트웨어학과 학부생",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "2동서울대학교 컴퓨터소프트웨어학과 학부생",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    add_para(doc, "jj81271000@gmail.com, m51187251@gmail.com",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)

    # 영어 제목
    add_para(doc,
             "DAAC: Disagreement-Aware Adaptive Consensus for Image Forgery Detection",
             size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)

    # 영어 저자/소속
    add_para(doc, "Jun-Jeong1, Moon-su hyeok2",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "1Dept. of Computer Software, Dong-Seoul University",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
    add_para(doc, "2Dept. of Computer Software, Dong-Seoul University",
             size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)

    # 요약 (테이블 박스)
    abs_tbl = doc.add_table(rows=1, cols=1)
    abs_tbl.style = 'Table Grid'
    cell = abs_tbl.rows[0].cells[0]

    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("요       약")
    set_font(r, 9, bold=True)

    abs_text = (
        "생성형 AI의 발전으로 이미지 위변조는 편집조작에서부터 Diffusion·GAN 기반 합성까지 다양해졌다. "
        "이에 따라 단일 원리에 의존하는 탐지기는 도메인 편향을 피하기 어려우므로 다중 전문 에이전트를 "
        "조합하는 접근이 자연스러운 해법이다. 본 연구에서는 에이전트 간 구조적 불일치를 활용하는 새로운 "
        "합의 방법 DAAC(Disagreement-Aware Adaptive Consensus for Image Forgery Detection)을 제시한다. "
        "4개 전문 에이전트(Frequency, Noise, FatFormer, Spatial)를 사용하여 고정 샘플셋(1,500장), "
        "10-seed 반복 평가한 결과, DAAC-GBM은 macro-F1 0.861로 COBRA(0.266) 대비 유의하게 우수했다"
        "(p=0.00195). 에이전트별 맹점—Frequency/Noise/Spatial의 AI-Gen F1=0.000, FatFormer의 "
        "Manipulated F1=0.000—이 DAAC에서는 상보적 불일치 신호로 전환되어 세 클래스 모두에서 균형 잡힌 "
        "성능을 달성한다. 6개 데이터 조합 60회 반복에서도 DAAC는 COBRA 대비 +0.493의 우위를 "
        "일관되게 유지했다(p=1.63e-11)."
    )
    p_abs = cell.add_paragraph()
    p_abs.paragraph_format.space_before = Pt(2)
    r = p_abs.add_run(abs_text)
    set_font(r, 9)

    # 빈 단락 제거
    empty = cell.paragraphs[0]
    empty._element.getparent().remove(empty._element)

    doc.add_paragraph()  # 요약 하단 여백

    # ── 연속 섹션 → 2단 전환 ────────────────────────────────────
    new_sec = doc.add_section(0)  # WD_SECTION.CONTINUOUS = 0
    new_sec.page_height = Cm(29.7)
    new_sec.page_width = Cm(21.0)
    new_sec.left_margin = Cm(2.5)
    new_sec.right_margin = Cm(2.5)
    new_sec.top_margin = Cm(2.5)
    new_sec.bottom_margin = Cm(2.5)
    set_two_columns(new_sec)

    # ── 2단 본문 ─────────────────────────────────────────────────

    # 1. 서론
    add_heading(doc, "1. 서론")
    for txt in [
        ("이미지 위변조 탐지는 조작 탐지와 생성형 이미지 기반 탐지를 동시에 만족해야 한다. "
         "실제 운영에서는 JPEG 압축 흔적, 센서 노이즈 불일치, 공간적 조작 흔적, 생성 모델 "
         "아티팩트가 혼재하므로, 단일 모델의 편향은 오탐/미탐으로 직결된다. 이에 이질적 탐지 "
         "원리를 가진 다수의 전문 에이전트를 조합하는 다중 에이전트 접근이 자연스러운 해법으로 "
         "제시된다."),
        ("그러나 다중 에이전트 합의(consensus) 단계가 병목이 된다. 기존 trust 기반 규칙형 합의"
         "(COBRA 계열: RoT/DRWA/AVGA)는 사전 고정된 글로벌 trust 점수를 기반으로 수식 기반 "
         "가중 집계를 수행한다. DRWA는 샘플별 분산으로 가중치를 조정하고 AVGA는 softmax "
         "temperature를 적응적으로 변경하지만, 이는 모두 handcrafted 규칙이며 에이전트들이 "
         "특정 클래스에 대해 가지는 구조적 맹점의 패턴을 데이터로부터 학습하지 못한다."),
        ('본 연구는 이러한 문제를 해결하기 위해 DAAC를 제안한다. 핵심 아이디어는 "누가 맞았는가"가 '
         '아닌 "어떻게 서로 다르게 틀리는가"를 학습 신호로 활용하는 것이다. 에이전트 판정 충돌로 '
         "구성된 43차원 메타 특징을 활용해서 GBM 분류기를 학습하면, 개별 에이전트들의 맹점이 "
         "오히려 중요한 근거가 됨을 정량적으로 확인하였다."),
    ]:
        add_para(doc, txt, indent=10)

    # 2. 관련 연구
    add_heading(doc, "2. 관련 연구")
    for txt in [
        ("다중 판정기 결합은 다수결/가중 다수결이 기본이다. COBRA[1]는 에이전트별 신뢰 점수"
         "(trust score) 기반 합의를 제공하며 RoT/DRWA/AVGA 세 알고리즘 변형을 포함한다. "
         "그러나 세 알고리즘 모두 사전 고정된 글로벌 trust 점수를 출발점으로 삼아 수식 기반 "
         "집계 함수는 학습되지 않는다. Oracle trust 조건에서도 F1이 소폭 하락하여, trust 품질보다 "
         "합의 방식 자체의 표현력 한계가 지배적임을 확인한다."),
        ("이미지 위변조 탐지 분야에서는 FatFormer[2]와 같은 생성 탐지기, TruFor[3]와 같은 "
         "다중 단서 기반 탐지기가 제안되었다. 그러나 단일 모델의 도메인 편향 문제는 여전하며, "
         "다중 에이전트 합의 단계에서 불일치 패턴 자체를 학습 신호로 활용한 연구는 드물다."),
    ]:
        add_para(doc, txt, indent=10)

    # 3. DAAC 방법론
    add_heading(doc, "3. DAAC 방법론")

    add_sub_heading(doc, "3.1 문제 정의")
    add_para(doc,
             "n개의 에이전트가 이미지 x에 대해 판정 vi∈C와 confidence ci∈[0,1]을 독립 산출한다. "
             "최종 클래스 C={authentic, manipulated, ai_generated}에 대한 합의 함수 f를 학습하는 "
             "것이 목표다. COBRA는 f를 사전 고정된 trust 점수 기반 handcrafted 수식으로 정의하지만, "
             "DAAC는 f를 데이터로부터 학습하며 불일치 구조를 포함한 43차원 메타 특징을 입력으로 사용한다.",
             indent=10)

    add_sub_heading(doc, "3.2 43차원 메타 특징 설계")
    add_para(doc,
             "DAAC 전체 특징(A5)은 총 43차원으로 구성된다: "
             "(1) Per-agent (20): verdict one-hot 4x4=16 + raw confidence 4; "
             "(2) Pairwise disagreement (18): 1[vi≠vj] 6쌍 + |ci-cj| 6쌍 + (ci+cj)·1[vi≠vj] 6쌍; "
             "(3) Aggregate (5): confidence variance, verdict entropy, max-min gap, "
             "unique verdict count, majority ratio. "
             "최상위 특징 중요도는 disagree_frequency_fatformer(56.5%)로, FatFormer와 나머지 "
             "에이전트 간 판정 불일치가 클래스 구분에 가장 유효한 신호임을 보인다.",
             indent=10)

    add_sub_heading(doc, "3.3 학습 모델")
    add_para(doc,
             "DAAC는 동일 43차원 특징셋에 대해 Logistic Regression(LR), Gradient Boosting Machine"
             "(GBM, XGBoost 백엔드), MLP를 학습한다. 제안 모델은 DAAC-GBM이며, "
             "경량 GBM으로 에이전트 inference 이후 추가 연산 부담이 미미하다.",
             indent=10)

    # 4. 실험 결과
    add_heading(doc, "4. 실험 결과")

    add_sub_heading(doc, "4.1 실험 설정")
    add_para(doc,
             "4개 전문 에이전트(Frequency/CAT-Net, Noise/MVSS, FatFormer, Spatial/Mesorch)로 "
             "구성된 MAIFS 플랫폼에서 실험하였다. 전체 1,500장(authentic/manipulated/ai_generated "
             "균형)을 고정 샘플셋으로 사용하고, train/val/test=0.6/0.2/0.2, seeds 300~309 "
             "(10회 반복)의 Protocol-P를 적용하였다. 통계 검정은 paired Wilcoxon을 사용한다.",
             indent=10)

    add_sub_heading(doc, "4.2 방법 간 비교")
    add_table(
        doc,
        headers=["방법", "Macro-F1\n(mean±std)", "p vs COBRA"],
        rows=[
            ["Majority Vote", "0.2774±0.0178", "0.0371"],
            ["Weighted MV", "0.2664±0.0128", "1.0000"],
            ["COBRA (drwa)", "0.2664±0.0090", "— (기준)"],
            ["Logistic Stack.\n(A4+LR)", "0.8500±0.0161", "0.00195**"],
            ["DAAC-LR", "0.8541±0.0087", "0.00195**"],
            ["DAAC-GBM (제안)", "0.8613±0.0156", "0.00195**"],
            ["DAAC-MLP", "0.8570±0.0122", "0.00195**"],
        ]
    )
    add_caption(doc, "<표 1> Macro-F1 비교 (10 seeds, **p<0.01)")
    add_para(doc,
             "DAAC 계열은 모든 10 seed에서 COBRA 대비 일관되게 우수하다. "
             "가장 단순한 Logistic Stacking(A4+LR)조차 +0.584의 절대 우위를 보이며, "
             "불일치 특징 학습 자체의 기여임을 시사한다.",
             indent=10)

    add_sub_heading(doc, "4.3 클래스별 성능")
    add_table(
        doc,
        headers=["방법", "Auth.", "Manip.", "AI-Gen."],
        rows=[
            ["Frequency", "0.804", "0.481", "0.000"],
            ["Noise", "0.558", "0.476", "0.000"],
            ["FatFormer", "0.542", "0.000", "0.473"],
            ["Spatial", "0.570", "0.355", "0.000"],
            ["COBRA", "0.573", "0.199", "0.027"],
            ["DAAC-GBM", "0.821", "0.800", "0.963"],
        ]
    )
    add_caption(doc, "<표 2> 클래스별 F1 (10-seed 평균)")
    add_para(doc,
             "Frequency/Noise/Spatial은 ai_generated F1=0.000, FatFormer는 manipulated F1=0.000의 "
             "구조적 맹점을 가진다. COBRA는 이 맹점들을 가중 합산하여 전파하는 반면, DAAC-GBM은 "
             "충돌 패턴을 학습하여 세 클래스 모두에서 균형 잡힌 성능을 달성한다.",
             indent=10)

    add_sub_heading(doc, "4.4 특징 Ablation")
    add_table(
        doc,
        headers=["특징 구성", "Dim", "Macro-F1"],
        rows=[
            ["A1: confidence only", "4", "0.7972±0.0149"],
            ["A2: verdict only", "16", "0.7596±0.0213"],
            ["A3: disagreement only", "23", "0.8534±0.0102"],
            ["A4: verdict+conf.", "20", "0.8506±0.0202"],
            ["A5: full DAAC (제안)", "43", "0.8613±0.0156"],
        ]
    )
    add_caption(doc, "<표 3> 특징 Ablation 결과 (GBM)")
    add_para(doc,
             "A3(불일치+집계, 23-dim)이 이미 0.8534로 높아, 불일치 구조 자체가 핵심 탐지 신호임을 "
             "확인한다. A3 > A4(0.8506)는 불일치 패턴이 개별 판정·confidence보다 강한 구분 신호임을 "
             "의미한다.",
             indent=10)

    add_sub_heading(doc, "4.5 다중 데이터 일반화")
    add_para(doc,
             "6개 데이터 조합 60회 반복(Protocol-M)에서 DAAC는 COBRA 대비 +0.493 우위를 "
             "일관되게 유지했다(Wilcoxon p=1.63e-11, sign pos/neg=60/0). "
             "단 한 번도 COBRA가 DAAC를 능가하지 못하였다.",
             indent=10)

    # 5. 결론
    add_heading(doc, "5. 결론")
    add_para(doc,
             "본 논문은 다중 에이전트 이미지 포렌식에서 에이전트 간 불일치 패턴을 명시적으로 학습하는 "
             "합의 메커니즘 DAAC를 제안하였다. 43차원 메타 특징 위의 GBM 분류기인 DAAC-GBM은 "
             "규칙형 합의(COBRA) 대비 macro-F1 0.861 vs 0.266(p=0.00195)의 유의한 우위를 보이며, "
             "6개 데이터 조합 60회 반복에서도 일관된 우위를 확인하였다(sign 60/0). "
             "에이전트 성능 개선 없이 합의 계층만 개선해도 +0.43~+0.55의 F1 향상이 가능하며, "
             "다중 에이전트 시스템에서 합의 방식 설계가 개별 에이전트 성능만큼 중요함을 시사한다.",
             indent=10)

    # 참고문헌
    add_heading(doc, "참고문헌")
    add_ref(doc,
            '[1] Z. Haider, et al. "Consensus-Based Reward Model for Better Alignment of Large '
            'Language Models" Scientific Reports, Vol. 15, Article 4004, 2025.')
    add_ref(doc,
            '[2] H. Liu, et al. "Learning to Detect AI-Generated Images via FatFormer" '
            'IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR), '
            'Seattle, 2024, pp. 10770-10780.')
    add_ref(doc,
            '[3] F. Guillaro, et al. "TruFor: Leveraging All-Round Clues for Trustworthy '
            'Image Forgery Detection and Localization" IEEE/CVF Conference on Computer '
            'Vision and Pattern Recognition (CVPR), Vancouver, 2023, pp. 20606-20615.')

    out = "/data/jj812_files/MAIFS/docs/research/DAAC_KIPS_Paper.docx"
    doc.save(out)
    print(f"완료: {out}")


if __name__ == "__main__":
    main()
